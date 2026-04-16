"""
Pre-process a DOCX textbook into plain text + extracted images.

Usage:
    python scripts/preprocess-docx.py "<path-to-docx>" "<output-dir>"

Example:
    python scripts/preprocess-docx.py ^
        "c:/Users/IAU/Documents/Claude Projects/e-books/dentistry/Oxford Handbook of Clinical Dentistry 7e.docx" ^
        "source-material/oxford-handbook"

Output:
    <output-dir>/
        full-text.txt        — all paragraphs, one per line, with style tags
        chapters.json        — chapter boundaries (paragraph indices)
        images/              — all images extracted, named by position
        image-map.json       — maps image filenames to paragraph index
"""

import json
import os
import sys
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def find_chapter_boundaries(paragraphs):
    chapters = []
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else ""
        if style in ("Para 086", "Heading 1", "Heading 2") and len(text) < 100:
            if any(kw in text.lower() for kw in [
                "dentistry", "examination", "surgery", "orthodontics",
                "paediatric", "medicine", "therapeutics", "analgesia",
                "materials", "ethics", "professionalism", "management",
                "syndromes", "addresses", "periodontology", "endodontics",
                "implants", "preventive", "community", "oral",
                "radiology", "replacing teeth", "repairing teeth",
            ]):
                chapters.append({"index": i, "title": text})
    return chapters


def extract_images(doc, out_dir):
    images_dir = out_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)

    image_map = []
    img_index = 0

    for i, p in enumerate(doc.paragraphs):
        for run in p.runs:
            if run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
                drawings = run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing")
            inline_imgs = run._element.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )
            for blip in inline_imgs:
                embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                if not embed:
                    continue
                rel = doc.part.rels.get(embed)
                if not rel or "image" not in rel.reltype:
                    continue

                img_data = rel.target_part.blob
                if len(img_data) < 2000:
                    continue

                ext = os.path.splitext(rel.target_ref)[1] or ".jpg"
                filename = f"para-{i:04d}-img-{img_index:03d}{ext}"
                with open(images_dir / filename, "wb") as f:
                    f.write(img_data)

                image_map.append({
                    "file": filename,
                    "paragraph": i,
                    "context": doc.paragraphs[i].text[:100] if doc.paragraphs[i].text else "",
                    "bytes": len(img_data),
                })
                img_index += 1

    return image_map


def main():
    if len(sys.argv) < 3:
        print("Usage: python scripts/preprocess-docx.py <docx-path> <output-dir>")
        sys.exit(1)

    docx_path = sys.argv[1]
    out_dir = Path(sys.argv[2])
    out_dir.mkdir(parents=True, exist_ok=True)

    print(f"Loading {docx_path}...")
    doc = Document(docx_path)
    print(f"  {len(doc.paragraphs)} paragraphs")

    # Extract full text
    print("Extracting text...")
    with open(out_dir / "full-text.txt", "w", encoding="utf-8") as f:
        for i, p in enumerate(doc.paragraphs):
            style = p.style.name if p.style else "Normal"
            text = p.text
            f.write(f"[{i}][{style}] {text}\n")

    # Find chapter boundaries
    print("Finding chapters...")
    chapters = find_chapter_boundaries(doc.paragraphs)
    with open(out_dir / "chapters.json", "w", encoding="utf-8") as f:
        json.dump(chapters, f, indent=2)
    print(f"  Found {len(chapters)} chapters")

    # Extract images
    print("Extracting images...")
    image_map = extract_images(doc, out_dir)
    with open(out_dir / "image-map.json", "w", encoding="utf-8") as f:
        json.dump(image_map, f, indent=2)
    print(f"  Extracted {len(image_map)} images")

    print("Done.")


if __name__ == "__main__":
    main()
